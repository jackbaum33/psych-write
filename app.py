from flask import Flask, render_template, request, jsonify, send_file
from docx import Document
import tempfile, os, platform, re
from openai import OpenAI
from docx2pdf import convert
import stripe
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

# Initialize
client = OpenAI()
app = Flask(__name__)

stripe.api_key = os.getenv("STRIPE_SECRET_KEY")
webhook_secret = os.getenv("STRIPE_WEBHOOK_SECRET")

stored_sessions = {}  # Stores form input by session ID
report_store = {}     # Stores generated report path by session ID

SECTIONS = {
    "referral": "Reason for Referral",
    "family_details": "Family Details",
    "birth_history": "Birth/Developmental History",
    "school_history": "School History",
    "previous_evals": "Previous Evaluations",
    "observations": "Behavioral Observations",
    "recommendations": "Recommendations"
}

def replace_placeholders(doc, replacements):
    for paragraph in doc.paragraphs:
        for key, val in replacements.items():
            if key in paragraph.text:
                for run in paragraph.runs:
                    if key in run.text:
                        run.text = run.text.replace(key, val)

def parse_numbered_paragraphs(response_text):
    pattern = r"###\s*\d+\.\s+[^\n]+\n(.+?)(?=\n###|\Z)"
    matches = re.findall(pattern, response_text, re.DOTALL)
    return [m.strip() for m in matches]

def batched_generate_paragraphs(section_prompts, appendix):
    combined_prompt = "".join(
        f"### {i+1}. {label}\n{content}\n" for i, (content, label) in enumerate(section_prompts)
    )
    prompt = (
        "You are a clinical psychologist drafting a neuropsychological report. "
        "For each section below, write a clear, professional paragraph.\n"
        "Incorporate relevant insights from the test appendix where applicable.\n"
        "Respond ONLY with the numbered paragraphs in the same order.\n\n"
        f"{combined_prompt}\n\nAppendix/Test Scores:\n{appendix}"
    )
    response = client.chat.completions.create(
        model="gpt-4o",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.7,
    )
    return parse_numbered_paragraphs(response.choices[0].message.content)

def generate_test_analysis(test_name, appendix):
    prompt = (
        f"Please write a paragraph analyzing the results and significance of the following neuropsychological test: {test_name}. "
        f"Use the appendix information below to guide the interpretation if relevant.\n\n"
        f"Appendix/Test Scores:\n{appendix}"
    )
    response = client.chat.completions.create(
        model="gpt-4o",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.7,
    )
    return response.choices[0].message.content.strip()

@app.route("/")
def index():
    return render_template("form.html", sections=SECTIONS)

@app.route("/store", methods=["POST"])
def store():
    session_id = request.form.get("session_id")
    form_data = request.form.to_dict(flat=False)
    stored_sessions[session_id] = form_data
    return "stored"

@app.route("/create-checkout-session", methods=["POST"])
def create_checkout():
    data = request.get_json()
    session_id = data.get("session_id")

    session = stripe.checkout.Session.create(
        payment_method_types=['card'],
        line_items=[{
            'price_data': {
                'currency': 'usd',
                'product_data': {'name': 'Neuropsych Report Generation'},
                'unit_amount': 1500,
            },
            'quantity': 1,
        }],
        mode='payment',
        success_url=f"http://psych-write.com/success?session_id={session_id}",
        cancel_url="http://psych-write.com/",
        metadata={"session_id": session_id}
    )
    return jsonify({"url": session.url})

@app.route("/webhook", methods=["POST"])
def stripe_webhook():
    payload = request.data
    sig_header = request.headers.get('stripe-signature')

    try:
        event = stripe.Webhook.construct_event(payload, sig_header, webhook_secret)
    except ValueError:
        return "Invalid payload", 400
    except stripe.error.SignatureVerificationError:
        return "Invalid signature", 400

    if event["type"] == "checkout.session.completed":
        session = event["data"]["object"]
        session_id = session["metadata"]["session_id"]

        data = stored_sessions.get(session_id)
        if not data:
            return "Missing stored data", 400

        report_path = generate_report(data)
        report_store[session_id] = report_path

    return "", 200

@app.route("/success")
def success():
    session_id = request.args.get("session_id")
    path = report_store.get(session_id)
    if path:
        return send_file(path, as_attachment=True)
    else:
        return "No report found for session.", 404

@app.route("/check-report")
def check_report():
    session_id = request.args.get("session_id")
    pdf_path = report_store.get(session_id)
    if pdf_path and os.path.exists(pdf_path):
        filename = "neuropsych_report.pdf" if pdf_path.endswith(".pdf") else "neuropsych_report.docx"
        return send_file(pdf_path, as_attachment=True, download_name=filename)
    return "Not ready", 404

def generate_report(data):
    doc = Document("doc_templates/report_template.docx")

    psychologist_name = data.get("psychologist_name", [""])[0]
    footer_information = data.get("footer_information", [psychologist_name])[0]
    appendix = data.get("appendix", [""])[0]

    section_prompts = []
    for key, label in SECTIONS.items():
        text = data.get(key, [""])[0].strip()
        if text:
            section_prompts.append((f"{label} Section\n{text}", f"{{{{{key}_paragraph}}}}"))

    paragraphs = batched_generate_paragraphs(section_prompts, appendix)

    test_sections = [ts for ts in data.get("test_types", []) if ts.strip()]
    test_texts = []
    for i, test_name in enumerate(test_sections, start=1):
        test_bullets = [b for b in data.get(f"test_{i}_bullets", []) if b.strip()]
        section_text = f"\n\n{test_name}:\n"
        if test_bullets:
            bullet_text = "\n".join(f"- {b}" for b in test_bullets)
            test_prompt = (
                f"Write a professional paragraph summarizing the following test and bullet points.\n"
                f"Test: {test_name}\nBullets:\n{bullet_text}\n\nAppendix:\n{appendix}"
            )
            test_response = client.chat.completions.create(
                model="gpt-4o",
                messages=[{"role": "user", "content": test_prompt}],
                temperature=0.7,
            )
            section_text += test_response.choices[0].message.content.strip() + "\n"
        analysis = generate_test_analysis(test_name, appendix)
        section_text += analysis
        test_texts.append(section_text)

    full_test_section = "\n\n".join(test_texts)
    test_list_string = ", ".join(test_sections)

    replacements = {
        "{{test_paragraphs}}": full_test_section,
        "{{test_list}}": test_list_string,
        "{{footer_information}}": footer_information,
        "{{appendix}}": appendix,
        "{{psychologist_name}}": psychologist_name,
    }

    for field in ["name", "dob", "age", "grade", "school", "eval_dates"]:
        replacements[f"{{{{{field}}}}}"] = data.get(field, [""])[0]

    for (_, placeholder), paragraph in zip(section_prompts, paragraphs):
        replacements[placeholder] = paragraph

    replace_placeholders(doc, replacements)
    section = doc.sections[-1]
    footer = section.footer
    footer.paragraphs[0].text = f"Report generated by: {psychologist_name}"

    tmpdir = tempfile.mkdtemp()
    docx_path = os.path.join(tmpdir, "report.docx")
    doc.save(docx_path)

    if "pdf" in data and data["pdf"]:
        sys = platform.system()
        pdf_path = os.path.join(tmpdir, "report.pdf")
        if sys in ["Windows", "Darwin"]:
            convert(docx_path, pdf_path)
            return pdf_path

    return docx_path

if __name__ == "__main__":
    app.run(debug=True)
